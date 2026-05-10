"""
Document upload, rich extraction (text + tables + images), RAG indexing, comparison.
Images and drawings are described by llava (vision LLM) and injected into the RAG text.
"""
from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from fastapi.responses import JSONResponse
import os, re, base64, json, httpx
from typing import List, Optional
from datetime import datetime
import aiofiles

from auth import get_current_user
from rag_pipeline import get_store, semantic_compare_rag

router     = APIRouter()
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

OLLAMA_BASE  = "http://localhost:11434"
VISION_MODEL = "llava"
FALLBACK_TEXT = "[Image could not be described — llava not available]"

# ── llava vision helper ───────────────────────────────────────────────────────

async def describe_image(base64_img: str, page_num: int, ext: str) -> str:
    prompt = (
        "You are analyzing a page from an academic or technical document. "
        "Describe what you see in detail. "
        "If it contains a TABLE: extract ALL rows and columns as plain text, preserving the data. "
        "If it contains a DIAGRAM, FLOWCHART, or DRAWING: explain every component, label, arrow, and relationship shown. "
        "If it contains a GRAPH or CHART: describe the axes, data trends, and key values. "
        "If it contains an EQUATION or FORMULA: write it out in plain text. "
        "Be thorough — this description will be used for question generation and search."
    )
    payload = {
        "model":  VISION_MODEL,
        "prompt": prompt,
        "images": [base64_img],
        "stream": False,
        "options": {"temperature": 0.1, "num_predict": 1000},
    }
    try:
        async with httpx.AsyncClient(timeout=120.0) as c:
            r = await c.post(f"{OLLAMA_BASE}/api/generate", json=payload)
            r.raise_for_status()
            description = r.json().get("response", "").strip()
            return description if description else FALLBACK_TEXT
    except httpx.ConnectError:
        return "[llava not available — run: ollama pull llava]"
    except httpx.HTTPStatusError as e:
        if e.response.status_code == 404:
            return "[llava model not found — run: ollama pull llava]"
        return f"[llava error: {e}]"
    except Exception as e:
        return f"[vision error: {e}]"


async def check_llava_available() -> bool:
    try:
        async with httpx.AsyncClient(timeout=5.0) as c:
            r = await c.get(f"{OLLAMA_BASE}/api/tags")
            models = [m["name"] for m in r.json().get("models", [])]
            return any("llava" in m.lower() for m in models)
    except Exception:
        return False


# ── Rich PDF extraction ───────────────────────────────────────────────────────

def extract_pdf_rich(filepath: str) -> dict:
    text_parts, tables, images = [], [], []
    page_count = 0

    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, 1):
                t = page.extract_text()
                if t and t.strip():
                    text_parts.append(t.strip())
                for tbl in page.extract_tables():
                    if not tbl: continue
                    rows = [[str(c or "").replace("\n", " ") for c in row] for row in tbl]
                    if not rows: continue
                    header = rows[0]
                    sep    = ["---"] * len(header)
                    body   = rows[1:] if len(rows) > 1 else []
                    md  = "| " + " | ".join(header) + " |\n"
                    md += "| " + " | ".join(sep)    + " |\n"
                    for row in body:
                        md += "| " + " | ".join(row) + " |\n"
                    tables.append({"page": page_num, "markdown": md, "rows": len(body), "cols": len(header)})
    except Exception as e:
        text_parts.append(f"[pdfplumber error: {e}]")

    try:
        import fitz
        doc = fitz.open(filepath)
        page_count = page_count or len(doc)
        for page_num, page in enumerate(doc, 1):
            for img_info in page.get_images(full=True):
                xref      = img_info[0]
                base_img  = doc.extract_image(xref)
                img_bytes = base_img["image"]
                ext       = base_img["ext"]
                if len(img_bytes) < 2048: continue
                b64 = base64.b64encode(img_bytes).decode("utf-8")
                images.append({
                    "page":        page_num,
                    "ext":         ext,
                    "size_kb":     round(len(img_bytes) / 1024, 1),
                    "data":        b64,
                    "description": None,
                })
        doc.close()
    except Exception:
        pass

    full_text = "\n\n".join(text_parts)
    if tables:
        full_text += "\n\n[TABLES EXTRACTED]\n"
        for t in tables:
            full_text += f"\n[Table on page {t['page']}]\n{t['markdown']}\n"

    return {"text": full_text, "tables": tables, "images": images, "page_count": page_count}


def extract_docx_rich(filepath: str) -> dict:
    text_parts, tables = [], []
    try:
        from docx import Document
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for tbl_idx, tbl in enumerate(doc.tables):
            rows = []
            for row in tbl.rows:
                rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
            if not rows: continue
            header = rows[0]
            sep    = ["---"] * len(header)
            body   = rows[1:]
            md  = "| " + " | ".join(header) + " |\n"
            md += "| " + " | ".join(sep)    + " |\n"
            for r in body:
                md += "| " + " | ".join(r) + " |\n"
            tables.append({"table_index": tbl_idx + 1, "markdown": md, "rows": len(body), "cols": len(header)})
        full_text = "\n\n".join(text_parts)
        if tables:
            full_text += "\n\n[TABLES]\n" + "\n".join(t["markdown"] for t in tables)
        return {"text": full_text, "tables": tables, "images": [], "page_count": None}
    except Exception as e:
        return {"text": f"Error extracting DOCX: {e}", "tables": [], "images": [], "page_count": None}


def extract_rich(filepath: str, filename: str) -> dict:
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":   return extract_pdf_rich(filepath)
    elif ext == "docx": return extract_docx_rich(filepath)
    elif ext in ["txt", "md"]:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return {"text": content, "tables": [], "images": [], "page_count": None}
    return {"text": "", "tables": [], "images": [], "page_count": None}


# ── Session ───────────────────────────────────────────────────────────────────

sessions: dict = {}

def get_session_docs(email: str) -> Optional[List[dict]]:
    return sessions.get(email, {}).get("docs")


# ── Routes ────────────────────────────────────────────────────────────────────

@router.post("/upload")
async def upload_documents(
    files: List[UploadFile] = File(...),
    current_user: dict = Depends(get_current_user),
):
    if len(files) > 2:
        raise HTTPException(400, "Maximum 2 documents allowed")

    email    = current_user["email"]
    user_dir = os.path.join(UPLOAD_DIR, email.replace("@", "_at_"))
    os.makedirs(user_dir, exist_ok=True)

    llava_ok = await check_llava_available()
    print(f"[upload] llava available: {llava_ok}")

    docs = []
    for idx, file in enumerate(files):
        ext = file.filename.lower().split(".")[-1]
        if ext not in ["pdf", "docx", "txt", "md"]:
            raise HTTPException(400, f"Unsupported file type: {ext}")

        filepath = os.path.join(user_dir, file.filename)
        async with aiofiles.open(filepath, "wb") as f:
            content = await file.read()
            await f.write(content)

        extracted = extract_rich(filepath, file.filename)

        # ── Vision pass ──────────────────────────────────────────────────────
        vision_descriptions = []
        if llava_ok and extracted["images"]:
            print(f"[upload] llava: describing {len(extracted['images'])} image(s) in {file.filename}")
            for img in extracted["images"]:
                print(f"[upload] → page {img['page']} ({img['size_kb']} KB)")
                desc = await describe_image(img["data"], img["page"], img["ext"])
                img["description"] = desc
                if desc and not desc.startswith("["):
                    vision_descriptions.append(
                        f"[Visual content on page {img['page']}]\n{desc}"
                    )

            if vision_descriptions:
                extracted["text"] += "\n\n[VISUAL CONTENT — described by AI]\n\n"
                extracted["text"] += "\n\n".join(vision_descriptions)
                print(f"[upload] added {len(vision_descriptions)} vision description(s) to RAG text")

        elif not llava_ok and extracted["images"]:
            print(f"[upload] skipping {len(extracted['images'])} image(s) — llava not found. Run: ollama pull llava")

        docs.append({
            "filename":         file.filename,
            "filepath":         filepath,
            "text":             extracted["text"],
            "tables":           extracted["tables"],
            "images":           extracted["images"],
            "page_count":       extracted["page_count"],
            "word_count":       len(extracted["text"].split()),
            "size":             len(content),
            "doc_index":        idx,
            "vision_used":      llava_ok and len(extracted["images"]) > 0,
            "images_described": len(vision_descriptions) if llava_ok else 0,
        })

    # RAG index — now includes vision descriptions
    store        = get_store(email)
    rag_docs     = [{"filename": d["filename"], "text": d["text"], "doc_index": d["doc_index"]} for d in docs]
    index_result = store.index_documents(rag_docs)

    result = {
        "documents": [
            {
                "filename":         d["filename"],
                "word_count":       d["word_count"],
                "size":             d["size"],
                "page_count":       d["page_count"],
                "table_count":      len(d["tables"]),
                "image_count":      len(d["images"]),
                "images_described": d["images_described"],
                "vision_used":      d["vision_used"],
                "tables": d["tables"][:5],
                "images": [
                    {
                        "page":        i["page"],
                        "ext":         i["ext"],
                        "size_kb":     i["size_kb"],
                        "description": i.get("description"),
                    }
                    for i in d["images"][:8]
                ],
            }
            for d in docs
        ],
        "rag":       {"chunks_indexed": index_result["chunks"], "status": "indexed"},
        "llava_used": llava_ok,
        "llava_tip":  None if llava_ok else "Run 'ollama pull llava' to enable image and diagram understanding.",
    }

    if len(docs) == 2:
        result["comparison"] = semantic_compare_rag(
            docs[0]["text"], docs[1]["text"],
            docs[0]["filename"], docs[1]["filename"],
        )

    sessions[email] = {"docs": docs, "uploaded_at": datetime.utcnow().isoformat()}
    return result


@router.get("/session")
async def get_session(current_user: dict = Depends(get_current_user)):
    email = current_user["email"]
    if email not in sessions:
        return {"documents": [], "has_docs": False, "rag_ready": False}
    store = get_store(email)
    meta  = store.get_metadata()
    docs_info = [
        {
            "filename":         d["filename"],
            "word_count":       d["word_count"],
            "table_count":      len(d["tables"]),
            "image_count":      len(d["images"]),
            "images_described": d.get("images_described", 0),
            "vision_used":      d.get("vision_used", False),
        }
        for d in sessions[email]["docs"]
    ]
    return {
        "documents": docs_info,
        "has_docs":  True,
        "rag_ready": meta.get("total_chunks", 0) > 0,
        "chunks":    meta.get("total_chunks", 0),
    }


@router.delete("/clear")
async def clear_session(current_user: dict = Depends(get_current_user)):
    email = current_user["email"]
    if email in sessions: del sessions[email]
    get_store(email).clear()
    return {"message": "Session and vector store cleared"}


@router.get("/rag-status")
async def rag_status(current_user: dict = Depends(get_current_user)):
    store = get_store(current_user["email"])
    meta  = store.get_metadata()
    return {
        "indexed":   meta.get("total_chunks", 0) > 0,
        "chunks":    meta.get("total_chunks", 0),
        "documents": meta.get("doc_names", []),
    }


@router.get("/llava-status")
async def llava_status():
    available = await check_llava_available()
    return {
        "available": available,
        "model":     VISION_MODEL,
        "tip":       None if available else "Run: ollama pull llava  to enable image/diagram understanding",
    }


@router.get("/images/{doc_index}")
async def get_doc_images(doc_index: int, current_user: dict = Depends(get_current_user)):
    email = current_user["email"]
    docs  = get_session_docs(email)
    if not docs: raise HTTPException(404, "No documents in session")
    if doc_index >= len(docs): raise HTTPException(404, "Document index out of range")
    doc = docs[doc_index]
    return {
        "filename": doc["filename"],
        "images": [
            {
                "page":        i["page"],
                "ext":         i["ext"],
                "size_kb":     i["size_kb"],
                "data":        i["data"],
                "description": i.get("description"),
            }
            for i in doc["images"]
        ],
    }
